"""Lay a conversation out as a Word document.

Deliberately domain-blind: this module has never heard of a load factor. The
browser sends prose, already-formatted stat tiles and already-rasterized plots,
and this file decides only where they sit on the page. That boundary is what
keeps the export honest — the figures in the .docx come from the same
frontend/src/charts/stats.js the screen renders from, so the two cannot drift,
and no KPI formatting is reimplemented in Python.
"""
import base64
import binascii
from io import BytesIO

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor

# The text column of the default template, so a plot spans it exactly.
PLOT_WIDTH = Inches(6.0)


def _provenance(payload):
    """When, and by which model — an exported answer outlives its session, and
    a reader months later needs the scope it was true in. Parts that are
    missing are dropped rather than printed empty."""
    parts = [payload.get("generated_at"), payload.get("model")]
    return " · ".join(part for part in parts if part)


def _png_bytes(raw):
    """Decode a capture, or None if it cannot be read.

    The payload is client-supplied, so these bytes are untrusted: a corrupt or
    truncated capture must cost that one picture and nothing else. Tolerates the
    'data:image/png;base64,' prefix canvas.toDataURL() produces, rather than
    pushing that stripping rule out to the browser.
    """
    if not raw:
        return None
    _, _, encoded = raw.rpartition(",")
    try:
        return base64.b64decode(encoded, validate=True)
    except (binascii.Error, ValueError):
        return None


def _add_stats(document, stats):
    """The headline numbers as a table: label, value, and the qualifier the tile
    carried beneath it. A tile with no note keeps its cell rather than shifting
    the columns."""
    if not stats:
        return
    table = document.add_table(rows=0, cols=3)
    table.style = "Table Grid"
    for stat in stats:
        cells = table.add_row().cells
        cells[0].text = str(stat.get("label", ""))
        cells[1].text = str(stat.get("value", ""))
        cells[2].text = str(stat.get("note") or "")


def _shade(cell, color):
    """Fill a table cell. python-docx has no shading API, so this writes the
    w:shd element the .docx schema wants directly onto the cell properties."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color.lstrip("#").upper())
    cell._tc.get_or_add_tcPr().append(shading)


def _add_bars(document, bars):
    """A part-to-whole bar, drawn as one shaded table row.

    Not every chart is an <svg>: TrafficMixChart is two CSS rectangles (see
    frontend/src/charts/TrafficMixChart.jsx), so the browser has nothing to
    photograph and sends the shares themselves. This stays domain-blind — it is
    handed labels, percentages and colors, and knows nothing about what they
    measure.
    """
    drawn = [bar for bar in bars if float(bar.get("pct") or 0) > 0]
    if not drawn:
        return

    table = document.add_table(rows=1, cols=len(drawn))
    table.autofit = False
    total = sum(float(bar["pct"]) for bar in drawn)
    for cell, bar in zip(table.rows[0].cells, drawn):
        # Width, not text length, is what makes this a bar: the segments have to
        # keep their proportions or the picture contradicts the numbers on it.
        cell.width = Emu(round(PLOT_WIDTH.emu * float(bar["pct"]) / total))
        _shade(cell, bar.get("color") or "#cccccc")
        run = cell.paragraphs[0].add_run(str(bar.get("label", "")))
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _add_note(document, text, italic=False):
    """A small-print line — caption or provenance."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(9)
    run.italic = italic
    return paragraph


def _add_chart(document, chart):
    """Numbers, then the plot they came from, then what it means — the reading
    order of the chart cards on screen (StatRow above, figcaption below)."""
    if chart.get("title"):
        document.add_heading(chart["title"], level=2)
    _add_stats(document, chart.get("stats") or [])

    _add_bars(document, chart.get("bars") or [])

    png = _png_bytes(chart.get("png_b64"))
    if png:
        try:
            document.add_picture(BytesIO(png), width=PLOT_WIDTH)
        except Exception:
            # Decodable but not an image python-docx recognises. Same rule as a
            # failed decode: lose the picture, never the finding beside it.
            pass

    if chart.get("caption"):
        _add_note(document, chart["caption"])
    if chart.get("source"):
        _add_note(document, f"Source: {chart['source']}", italic=True)


def build_document(payload):
    """A document description in, .docx bytes out."""
    document = Document()
    document.add_heading(payload.get("title", "Conversation"), level=0)
    document.add_paragraph(_provenance(payload))

    for index, turn in enumerate(payload.get("turns") or [], start=1):
        document.add_heading(f"{index}. {turn.get('question', '')}", level=1)
        document.add_paragraph(turn.get("answer", ""))
        for chart in (turn.get("charts") or []):
            _add_chart(document, chart)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
